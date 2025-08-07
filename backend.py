from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
import google.generativeai as genai
import os
import json
import asyncio
import logging
from typing import Optional, Dict, Any
import uuid
from datetime import datetime
import pdfplumber
import docx
from io import BytesIO
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(title="Hey Sera API", description="AI Assistant for Policy Document Analysis")

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # Add your frontend URL
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configure Gemini API
GEMINI_API_KEY = "AIzaSyDSinmAUdLRA2yZb--hxN9i_umbQoSwpOY"  # Consider moving this to an environment variable
genai.configure(api_key=GEMINI_API_KEY)

# Initialize Gemini model
model = genai.GenerativeModel('gemini-1.5-flash')

# Local storage configuration
MEMORY_DIR = Path("memory")
CHAT_SESSIONS_FILE = MEMORY_DIR / "chat_sessions.json"
UPLOADED_DOCUMENTS_FILE = MEMORY_DIR / "uploaded_documents.json"

# Create memory directory if it doesn't exist
MEMORY_DIR.mkdir(exist_ok=True)

# Initialize storage files if they don't exist
def init_storage_files():
    """Initialize JSON storage files if they don't exist"""
    if not CHAT_SESSIONS_FILE.exists():
        with open(CHAT_SESSIONS_FILE, 'w', encoding='utf-8') as f:
            json.dump({}, f, indent=2)
    
    if not UPLOADED_DOCUMENTS_FILE.exists():
        with open(UPLOADED_DOCUMENTS_FILE, 'w', encoding='utf-8') as f:
            json.dump({}, f, indent=2)

def load_chat_sessions() -> Dict[str, Dict[str, Any]]:
    """Load chat sessions from JSON file"""
    try:
        if CHAT_SESSIONS_FILE.exists():
            with open(CHAT_SESSIONS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    except Exception as e:
        logger.error(f"Error loading chat sessions: {e}")
        return {}

def save_chat_sessions(sessions: Dict[str, Dict[str, Any]]):
    """Save chat sessions to JSON file"""
    try:
        with open(CHAT_SESSIONS_FILE, 'w', encoding='utf-8') as f:
            json.dump(sessions, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logger.error(f"Error saving chat sessions: {e}")

def load_uploaded_documents() -> Dict[str, Dict[str, Any]]:
    """Load uploaded documents from JSON file"""
    try:
        if UPLOADED_DOCUMENTS_FILE.exists():
            with open(UPLOADED_DOCUMENTS_FILE, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    except Exception as e:
        logger.error(f"Error loading uploaded documents: {e}")
        return {}

def save_uploaded_documents(documents: Dict[str, Dict[str, Any]]):
    """Save uploaded documents to JSON file"""
    try:
        with open(UPLOADED_DOCUMENTS_FILE, 'w', encoding='utf-8') as f:
            json.dump(documents, f, indent=2, ensure_ascii=False)
    except Exception as e:
        logger.error(f"Error saving uploaded documents: {e}")

def backup_data():
    """Create a backup of current data with timestamp"""
    try:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        backup_dir = MEMORY_DIR / "backups"
        backup_dir.mkdir(exist_ok=True)
        
        # Backup chat sessions
        if CHAT_SESSIONS_FILE.exists():
            backup_chat_file = backup_dir / f"chat_sessions_{timestamp}.json"
            with open(CHAT_SESSIONS_FILE, 'r', encoding='utf-8') as src:
                with open(backup_chat_file, 'w', encoding='utf-8') as dst:
                    dst.write(src.read())
        
        # Backup documents
        if UPLOADED_DOCUMENTS_FILE.exists():
            backup_docs_file = backup_dir / f"uploaded_documents_{timestamp}.json"
            with open(UPLOADED_DOCUMENTS_FILE, 'r', encoding='utf-8') as src:
                with open(backup_docs_file, 'w', encoding='utf-8') as dst:
                    dst.write(src.read())
                    
        logger.info(f"Data backup created with timestamp: {timestamp}")
        return timestamp
    except Exception as e:
        logger.error(f"Error creating backup: {e}")
        return None

# Initialize storage on startup
init_storage_files()

# Pydantic models for request/response
class ChatMessage(BaseModel):
    message: str
    chatId: Optional[str] = None

class ChatResponse(BaseModel):
    response: str
    chatId: str
    timestamp: str

class DocumentUpload(BaseModel):
    filename: str
    content: str
    chatId: Optional[str] = None

# System prompt for Sera
SYSTEM_PROMPT = """
You are Sera, an AI assistant specialized in policy document analysis. Your role is to:

1. Analyze policy documents with precision and clarity
2. Answer questions about policy content, implications, and recommendations
3. Provide summaries, key points, and actionable insights
4. Help users understand complex policy language in simple terms
5. Identify potential issues, gaps, or contradictions in policies
6. Suggest improvements or alternatives when appropriate

Always maintain a professional yet friendly tone. Be concise but thorough in your responses. 
When analyzing documents, focus on:
- Key objectives and goals
- Implementation requirements
- Potential challenges or risks
- Stakeholder impacts
- Compliance requirements
- Financial implications

If you don't have enough context or information, ask clarifying questions.
"""

def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file using pdfplumber for better accuracy"""
    try:
        text = ""
        with pdfplumber.open(BytesIO(file_content)) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += f"\n--- Page {page_num} ---\n"
                        text += page_text + "\n"
                    
                    # Also extract text from tables if present
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, 1):
                            text += f"\n--- Table {table_num} on Page {page_num} ---\n"
                            for row in table:
                                if row and any(cell for cell in row if cell):  # Skip empty rows
                                    text += " | ".join(str(cell) if cell else "" for cell in row) + "\n"
                            text += "\n"
                            
                except Exception as page_error:
                    logger.warning(f"Error extracting from page {page_num}: {page_error}")
                    continue
                    
        # Clean up the text
        text = text.strip()
        if not text:
            logger.warning("No text extracted from PDF")
            return ""
            
        logger.info(f"Successfully extracted {len(text)} characters from PDF")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting PDF text with pdfplumber: {e}")
        # Fallback to basic text extraction if pdfplumber fails
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        except Exception as fallback_error:
            logger.error(f"Fallback PDF extraction also failed: {fallback_error}")
            return ""

def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(BytesIO(file_content))
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            text += "\n--- Table ---\n"
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_text.append(cell_text)
                if row_text:
                    text += " | ".join(row_text) + "\n"
            text += "\n"
        
        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text
        
    except Exception as e:
        logger.error(f"Error extracting DOCX text: {e}")
        return ""

def get_or_create_chat_session(chat_id: Optional[str] = None) -> str:
    """Get existing chat session or create a new one"""
    chat_sessions = load_chat_sessions()
    
    if not chat_id or chat_id not in chat_sessions:
        chat_id = str(uuid.uuid4())
        chat_sessions[chat_id] = {
            "id": chat_id,
            "messages": [],
            "documents": [],
            "created_at": datetime.now().isoformat(),
            "last_updated": datetime.now().isoformat()
        }
        save_chat_sessions(chat_sessions)
    
    return chat_id

async def generate_response(prompt: str, chat_id: str, context: str = "") -> str:
    """Generate response using Gemini API"""
    try:
        chat_sessions = load_chat_sessions()
        
        # Prepare the full prompt with context
        full_prompt = f"{SYSTEM_PROMPT}\n\n"
        
        # Add document context if available
        if context:
            full_prompt += f"Document Context:\n{context}\n\n"
        
        # Add chat history for context
        if chat_id in chat_sessions:
            recent_messages = chat_sessions[chat_id]["messages"][-10:]  # Last 10 messages
            if recent_messages:
                full_prompt += "Recent conversation:\n"
                for msg in recent_messages:
                    role = "User" if msg["role"] == "user" else "Sera"
                    full_prompt += f"{role}: {msg['content']}\n"
                full_prompt += "\n"
        
        full_prompt += f"User: {prompt}\nSera:"
        
        # Generate response
        response = await asyncio.to_thread(model.generate_content, full_prompt)
        return response.text
    
    except Exception as e:
        logger.error(f"Error generating response: {e}")
        return "I apologize, but I encountered an error processing your request. Please try again."

@app.post("/api/chat", response_model=ChatResponse)
async def chat_endpoint(message: ChatMessage):
    """Main chat endpoint"""
    try:
        # Get or create chat session
        chat_id = get_or_create_chat_session(message.chatId)
        
        # Load current data
        chat_sessions = load_chat_sessions()
        uploaded_documents = load_uploaded_documents()
        
        # Get document context if available
        context = ""
        if chat_id in chat_sessions and chat_sessions[chat_id]["documents"]:
            # Combine all document content for context
            doc_texts = []
            for doc_id in chat_sessions[chat_id]["documents"]:
                if doc_id in uploaded_documents:
                    doc_texts.append(uploaded_documents[doc_id]["content"])
            context = "\n\n".join(doc_texts)
        
        # Generate response
        response_text = await generate_response(message.message, chat_id, context)
        
        # Store messages in session
        chat_sessions[chat_id]["messages"].extend([
            {"role": "user", "content": message.message, "timestamp": datetime.now().isoformat()},
            {"role": "assistant", "content": response_text, "timestamp": datetime.now().isoformat()}
        ])
        chat_sessions[chat_id]["last_updated"] = datetime.now().isoformat()
        
        # Save updated sessions
        save_chat_sessions(chat_sessions)
        
        return ChatResponse(
            response=response_text,
            chatId=chat_id,
            timestamp=datetime.now().isoformat()
        )
    
    except Exception as e:
        logger.error(f"Chat endpoint error: {e}")
        raise HTTPException(status_code=500, detail="Internal server error")

@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...), chatId: Optional[str] = None):
    """Upload and process document"""
    try:
        # Validate file type
        allowed_types = [".pdf", ".docx", ".txt"]
        file_extension = os.path.splitext(file.filename)[1].lower()
        if file_extension not in allowed_types:
            raise HTTPException(
                status_code=400, 
                detail=f"Unsupported file type: {file_extension}. Supported types: {', '.join(allowed_types)}"
            )
        
        # Validate file size (10MB limit)
        file_content = await file.read()
        max_size = 10 * 1024 * 1024  # 10MB
        if len(file_content) > max_size:
            raise HTTPException(status_code=400, detail="File too large. Maximum size is 10MB.")
        
        logger.info(f"Processing file: {file.filename} ({len(file_content)} bytes)")
        
        # Extract text based on file type
        if file_extension == ".pdf":
            text_content = extract_text_from_pdf(file_content)
        elif file_extension == ".docx":
            text_content = extract_text_from_docx(file_content)
        else:  # .txt
            try:
                text_content = file_content.decode('utf-8')
            except UnicodeDecodeError:
                try:
                    text_content = file_content.decode('latin-1')
                except UnicodeDecodeError:
                    text_content = file_content.decode('utf-8', errors='ignore')
        
        if not text_content or not text_content.strip():
            raise HTTPException(
                status_code=400, 
                detail="Could not extract text from document. The file might be empty or corrupted."
            )
        
        # Load current data
        uploaded_documents = load_uploaded_documents()
        chat_sessions = load_chat_sessions()
        
        # Store document
        doc_id = str(uuid.uuid4())
        uploaded_documents[doc_id] = {
            "id": doc_id,
            "filename": file.filename,
            "content": text_content,
            "uploaded_at": datetime.now().isoformat(),
            "file_type": file_extension,
            "file_size": len(file_content),
            "text_length": len(text_content)
        }
        
        # Save documents
        save_uploaded_documents(uploaded_documents)
        
        # Associate with chat session
        chat_id = get_or_create_chat_session(chatId)
        chat_sessions = load_chat_sessions()  # Reload to get latest data
        chat_sessions[chat_id]["documents"].append(doc_id)
        chat_sessions[chat_id]["last_updated"] = datetime.now().isoformat()
        save_chat_sessions(chat_sessions)
        
        # Generate initial analysis
        analysis_prompt = f"I've uploaded a document titled '{file.filename}'. Please provide a brief summary and key insights from this policy document."
        analysis = await generate_response(analysis_prompt, chat_id, text_content)
        
        logger.info(f"Successfully processed document: {file.filename}")
        
        return {
            "success": True,
            "documentId": doc_id,
            "chatId": chat_id,
            "filename": file.filename,
            "fileSize": len(file_content),
            "textLength": len(text_content),
            "analysis": analysis
        }
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.get("/api/chat/{chat_id}/history")
async def get_chat_history(chat_id: str):
    """Get chat history for a session"""
    chat_sessions = load_chat_sessions()
    if chat_id not in chat_sessions:
        raise HTTPException(status_code=404, detail="Chat session not found")
    
    return chat_sessions[chat_id]

@app.get("/api/chat/{chat_id}/documents")
async def get_chat_documents(chat_id: str):
    """Get documents associated with a chat session"""
    chat_sessions = load_chat_sessions()
    uploaded_documents = load_uploaded_documents()
    
    if chat_id not in chat_sessions:
        raise HTTPException(status_code=404, detail="Chat session not found")
    
    documents = []
    for doc_id in chat_sessions[chat_id]["documents"]:
        if doc_id in uploaded_documents:
            doc = uploaded_documents[doc_id].copy()
            # Don't return the full content, just metadata
            doc.pop("content", None)
            documents.append(doc)
    
    return {"documents": documents}

@app.get("/api/chats")
async def get_all_chats():
    """Get all chat sessions (metadata only)"""
    chat_sessions = load_chat_sessions()
    
    # Return only metadata, not full message content
    chats = []
    for chat_id, session in chat_sessions.items():
        chat_metadata = {
            "id": session["id"],
            "created_at": session["created_at"],
            "last_updated": session.get("last_updated", session["created_at"]),
            "message_count": len(session["messages"]),
            "document_count": len(session["documents"])
        }
        
        # Add last message preview if available
        if session["messages"]:
            last_user_message = None
            for msg in reversed(session["messages"]):
                if msg["role"] == "user":
                    last_user_message = msg["content"][:100] + "..." if len(msg["content"]) > 100 else msg["content"]
                    break
            chat_metadata["last_message_preview"] = last_user_message
        
        chats.append(chat_metadata)
    
    # Sort by last_updated, most recent first
    chats.sort(key=lambda x: x["last_updated"], reverse=True)
    
    return {"chats": chats}

@app.delete("/api/chat/{chat_id}")
async def delete_chat(chat_id: str):
    """Delete a chat session"""
    chat_sessions = load_chat_sessions()
    uploaded_documents = load_uploaded_documents()
    
    if chat_id in chat_sessions:
        # Also remove associated documents
        for doc_id in chat_sessions[chat_id]["documents"]:
            uploaded_documents.pop(doc_id, None)
        
        chat_sessions.pop(chat_id, None)
        
        # Save updated data
        save_chat_sessions(chat_sessions)
        save_uploaded_documents(uploaded_documents)
        
        return {"success": True, "message": "Chat session deleted"}
    
    raise HTTPException(status_code=404, detail="Chat session not found")

@app.post("/api/backup")
async def create_backup():
    """Create a backup of all data"""
    timestamp = backup_data()
    if timestamp:
        return {"success": True, "message": f"Backup created successfully", "timestamp": timestamp}
    else:
        raise HTTPException(status_code=500, detail="Failed to create backup")

@app.get("/api/health")
async def health_check():
    """Health check endpoint"""
    chat_sessions = load_chat_sessions()
    uploaded_documents = load_uploaded_documents()
    
    return {
        "status": "healthy",
        "service": "Hey Sera API",
        "timestamp": datetime.now().isoformat(),
        "active_sessions": len(chat_sessions),
        "uploaded_documents": len(uploaded_documents),
        "storage": {
            "chat_sessions_file": str(CHAT_SESSIONS_FILE),
            "documents_file": str(UPLOADED_DOCUMENTS_FILE),
            "memory_dir": str(MEMORY_DIR)
        }
    }

@app.get("/api/stats")
async def get_stats():
    """Get system statistics"""
    chat_sessions = load_chat_sessions()
    uploaded_documents = load_uploaded_documents()
    
    total_messages = sum(len(session["messages"]) for session in chat_sessions.values())
    total_documents = len(uploaded_documents)
    total_text_length = sum(doc.get("text_length", 0) for doc in uploaded_documents.values())
    
    return {
        "total_chats": len(chat_sessions),
        "total_messages": total_messages,
        "total_documents": total_documents,
        "total_text_length": total_text_length,
        "storage_files": {
            "chat_sessions_exists": CHAT_SESSIONS_FILE.exists(),
            "documents_exists": UPLOADED_DOCUMENTS_FILE.exists(),
            "chat_sessions_size": CHAT_SESSIONS_FILE.stat().st_size if CHAT_SESSIONS_FILE.exists() else 0,
            "documents_size": UPLOADED_DOCUMENTS_FILE.stat().st_size if UPLOADED_DOCUMENTS_FILE.exists() else 0
        }
    }

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)